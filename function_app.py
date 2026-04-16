import azure.functions as func
import logging
import os
import io
import re
import ast
from urllib.parse import urlparse, unquote

import pandas as pd
import requests
from docx import Document
from docx.shared import Inches, Pt
from azure.storage.blob import BlobServiceClient
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

app = func.FunctionApp()

SHEET_NAME = "Merged"
ID_COLUMN = "ID Case"
CATALOG_COL = "Catalog Image Path"
SCREEN_COL = "Screencapture Path"

INPUT_CONTAINER = "input"
OUTPUT_CONTAINER = "output"
MERGED_EXCEL_BLOB = "merged.xlsx"

STORAGE_CONN_SETTING = "AzureWebJobsStorage"


def set_font(run):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def apply_normal_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def parse_urls_from_cell(value):
    if value is None or pd.isna(value):
        return []

    if isinstance(value, (list, tuple)):
        return [str(x).strip() for x in value if str(x).strip()]

    text = str(value).strip()
    if not text:
        return []

    try:
        maybe_list = ast.literal_eval(text)
        if isinstance(maybe_list, (list, tuple)):
            return [str(x).strip() for x in maybe_list if str(x).strip()]
    except Exception:
        pass

    for sep in [";", "\n", ","]:
        if sep in text:
            return [p.strip() for p in text.split(sep) if p.strip()]

    return [text]


def extract_case_id_from_filename(name: str) -> str:
    base = os.path.splitext(os.path.basename(name))[0]
    nums = re.findall(r"\d+", base)
    if not nums:
        raise ValueError(f"No numeric Case ID found in filename: {name}")
    nums.sort(key=len, reverse=True)
    return nums[0]


def download_bytes(url: str, timeout_s: int = 20):
    try:
        r = requests.get(url, timeout=timeout_s)
        r.raise_for_status()
        return r.content
    except Exception as ex:
        logging.error(f"Download failed: {url} -> {ex}")
        return None


def parse_blob_path_from_event(event: func.EventGridEvent) -> tuple[str, str]:
    data = event.get_json() or {}
    url = data.get("url")
    if not url:
        raise ValueError("EventGrid payload missing data.url")

    parsed = urlparse(url)
    path = unquote(parsed.path).lstrip("/")
    parts = path.split("/", 1)
    if len(parts) != 2:
        raise ValueError(f"Unexpected blob url path: {parsed.path}")

    container, blob_path = parts[0], parts[1]
    return container, blob_path


def process_docx_blob(blob_service_client: BlobServiceClient, blob_path: str):
    logging.info(f"Processing blob path: {blob_path}")

    case_id = extract_case_id_from_filename(blob_path)
    logging.info(f"Case ID: {case_id}")

    excel_blob = blob_service_client.get_blob_client(
        container=INPUT_CONTAINER,
        blob=MERGED_EXCEL_BLOB
    )
    excel_bytes = excel_blob.download_blob().readall()
    df = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=SHEET_NAME)

    df[ID_COLUMN] = df[ID_COLUMN].astype(str).str.strip()
    hit = df.loc[df[ID_COLUMN] == str(case_id)]

    if hit.empty:
        logging.error(f"No matching Case ID in Excel: {case_id}")
        return

    row = hit.iloc[0]
    catalog_urls = parse_urls_from_cell(row[CATALOG_COL])
    screen_urls = parse_urls_from_cell(row[SCREEN_COL])

    docx_blob = blob_service_client.get_blob_client(
        container=INPUT_CONTAINER,
        blob=blob_path
    )
    docx_bytes = docx_blob.download_blob().readall()

    doc = Document(io.BytesIO(docx_bytes))
    apply_normal_style(doc)

    def add_section(title, subtitle, urls):
        if not urls:
            return
        if len(urls) > 1:
            if title == "3. számú melléklet / Annex 3":
                subtitle = "Fotók / Images"
            elif title == "4. számú melléklet / Annex 4":
                subtitle = "Képernyőfotók a Fotók jogsértő felhasználásáról / Print screens of the infringing use of the Images"

        doc.add_page_break()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = title.split('/')
        run = p_title.add_run(parts[0].strip())
        run.bold = True
        set_font(run)

        run = p_title.add_run(" / ")
        set_font(run)

        run = p_title.add_run(parts[1].strip())
        run.bold = True
        run.italic = True
        set_font(run)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_parts = subtitle.split('/')
        run = p_sub.add_run(sub_parts[0].strip())
        run.bold = True
        set_font(run)

        run = p_sub.add_run(" / ")
        set_font(run)

        run = p_sub.add_run(sub_parts[1].strip())
        run.bold = True
        run.italic = True
        set_font(run)

        for url in urls:
            img = download_bytes(url)
            if not img:
                p = doc.add_paragraph("[IMAGE ERROR]")
                for r in p.runs:
                    set_font(r)
                continue
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                run.add_picture(io.BytesIO(img), width=Inches(5.5))
            except Exception:
                p = doc.add_paragraph("[INSERT ERROR]")
                for r in p.runs:
                    set_font(r)

    add_section("3. számú melléklet / Annex 3", "Fotó / Image", catalog_urls)
    add_section("4. számú melléklet / Annex 4", "Képernyőfotó a Fotó jogsértő felhasználásáról / Print screen of the infringing use of the Image", screen_urls)

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    output_name = f"docs_with_images/{case_id}_with_images.docx"

    output_blob = blob_service_client.get_blob_client(
        container=OUTPUT_CONTAINER,
        blob=output_name
    )
    output_blob.upload_blob(output_stream, overwrite=True)

    logging.info(f"Finished processing {case_id}")


@app.event_grid_trigger(arg_name="event")
def ProcessDocxEvent(event: func.EventGridEvent):

    container, blob_path = parse_blob_path_from_event(event)
    logging.info(f"EventGrid received container={container}, blob={blob_path}")

    if container != INPUT_CONTAINER or not blob_path.startswith("docs/"):
        logging.info("Event ignored (not input/docs)")
        return

    conn_str = os.environ[STORAGE_CONN_SETTING]
    blob_service_client = BlobServiceClient.from_connection_string(conn_str)

    process_docx_blob(blob_service_client, blob_path)